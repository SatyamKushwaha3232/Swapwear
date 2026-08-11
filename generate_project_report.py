from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION_START
from pathlib import Path
from datetime import date

OUT = Path('reports')
OUT.mkdir(exist_ok=True)
FILE = OUT / 'SwapWear_Internship_Project_Report.docx'

NAVY='26364A'; TEAL='177E89'; MINT='EAF5F3'; PALE='F3F6F8'; GRAY='5B6673'; WHITE='FFFFFF'

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def set_cell(cell, text, bold=False, color='000000', size=9.5):
    cell.text=''; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(str(text)); r.bold=bold; r.font.size=Pt(size); r.font.name='Aptos'; r._element.rPr.rFonts.set(qn('w:ascii'),'Aptos'); r.font.color.rgb=RGBColor.from_string(color); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
def border(cell, color='D8E0E6'):
    tcPr=cell._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        e=OxmlElement(f'w:{side}'); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4'); e.set(qn('w:color'),color); b.append(e)
    tcPr.append(b)
def table(doc, headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; shade(c,NAVY); set_cell(c,h,True,WHITE,9); border(c,NAVY)
        if widths: c.width=Inches(widths[i])
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,val in enumerate(row):
            c=cells[i]; shade(c,'FFFFFF' if ri%2==0 else PALE); set_cell(c,val,False,'202B36',9); border(c)
            if widths: c.width=Inches(widths[i])
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t
def para(doc, text='', bold=False, size=10.5, color='202B36', align=None, after=7, before=0, italic=False):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before); p.paragraph_format.line_spacing=1.2
    if align: p.alignment=align
    r=p.add_run(text); r.bold=bold; r.italic=italic; r.font.name='Aptos'; r._element.rPr.rFonts.set(qn('w:ascii'),'Aptos'); r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color)
    return p
def bullet(doc,text):
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.15
    r=p.add_run(text); r.font.name='Aptos'; r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string('202B36')
def heading(doc,text,level=1):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(15 if level==1 else 10); p.paragraph_format.space_after=Pt(6)
    r=p.add_run(text); r.bold=True; r.font.name='Aptos Display'; r._element.rPr.rFonts.set(qn('w:ascii'),'Aptos Display'); r.font.size=Pt(16 if level==1 else 12.5); r.font.color.rgb=RGBColor.from_string(NAVY if level==1 else TEAL)
    return p
def pagebreak(doc): doc.add_page_break()
def footer(section):
    p=section.footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('SwapWear Internship Project Report  |  2026'); r.font.name='Aptos'; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(GRAY)
def setup(doc):
    s=doc.sections[0]; s.top_margin=Inches(.78); s.bottom_margin=Inches(.7); s.left_margin=Inches(.8); s.right_margin=Inches(.8); footer(s)
    styles=doc.styles; normal=styles['Normal']; normal.font.name='Aptos'; normal._element.rPr.rFonts.set(qn('w:ascii'),'Aptos'); normal.font.size=Pt(10.5)

doc=Document(); setup(doc)
# Cover
para(doc,'INTERNSHIP PROJECT REPORT',True,11,TEAL,WD_ALIGN_PARAGRAPH.CENTER,after=20,before=50)
para(doc,'SwapWear',True,31,NAVY,WD_ALIGN_PARAGRAPH.CENTER,after=7)
para(doc,'A Premium Peer-to-Peer Clothing Swap Marketplace',False,15,GRAY,WD_ALIGN_PARAGRAPH.CENTER,after=26)
table(doc,['Project profile','Details'],[
 ['Project type','Full-stack web application'],['Internship organisation','________________________________'],['Internship duration','________________________________'],['Industry mentor','________________________________'],['Repository','https://github.com/SatyamKushwaha3232/Swapwear'],['Live deployment','https://swapwear-iota.vercel.app/'],['Frontend','React, Vite, Tailwind CSS'],['Backend','Node.js, Express, Socket.IO'],['Data layer','PostgreSQL with Prisma ORM'],['Prepared by','Satyam Kushwaha'],['Submission date','7 August 2026']], [2.0,4.5])
para(doc,'Submitted as an internship project report.',False,10,GRAY,WD_ALIGN_PARAGRAPH.CENTER,after=18)
para(doc,'Institution / Department: ________________________________',False,10,GRAY,WD_ALIGN_PARAGRAPH.CENTER,after=4)
para(doc,'Faculty guide: ___________________________________________',False,10,GRAY,WD_ALIGN_PARAGRAPH.CENTER,after=4)
pagebreak(doc)

heading(doc,'Certificate',1); para(doc,'This is to certify that the project entitled “SwapWear: A Premium Peer-to-Peer Clothing Swap Marketplace” has been prepared by Satyam Kushwaha as part of the required academic/project work. The work presented in this report is based on the implemented application and its accompanying source repository.');
para(doc,'Guide signature: ____________________    Head of Department: ____________________',False,10,GRAY,after=14)
heading(doc,'Declaration',1); para(doc,'I declare that this report describes my original project work. Wherever project tools, frameworks, libraries, or external services are referenced, they are identified in the relevant sections.');
para(doc,'Student signature: ____________________    Date: ____________________',False,10,GRAY,after=10)
heading(doc,'Acknowledgement',1); para(doc,'I would like to thank my faculty guide, department, peers, and all contributors who supported the planning, implementation, testing, and review of SwapWear. Their feedback helped shape the application into a complete marketplace workflow.');
heading(doc,'Internship Details',1); para(doc,'Organisation name: ________________________________. Internship duration: ________________________________. Industry mentor: ________________________________. This report should be submitted after completing these organisation-specific details.');
heading(doc,'Roles and Responsibilities',2)
for x in ['Analysed marketplace requirements and converted them into user-facing application modules.','Implemented responsive React pages, reusable components, protected routes, and service layers.','Developed backend APIs and business logic using Express, Prisma, and PostgreSQL.','Designed data models, swap lifecycle controls, validation, testing workflows, and deployment documentation.']: bullet(doc,x)
pagebreak(doc)

heading(doc,'Abstract',1); para(doc,'SwapWear was developed as an internship project to create a responsive full-stack platform for exchanging pre-owned clothing through structured peer-to-peer swaps. The work involved requirement analysis, interface development, API design, relational database modelling, feature implementation, testing, and deployment preparation. It helps users list wearable items, discover suitable alternatives, submit swap requests, communicate with other members, and track each transaction to completion.');
para(doc,'The system combines a React and Vite frontend with a Node.js/Express backend. PostgreSQL and Prisma provide the relational data model required for users, listings, swaps, deliveries, payments, chats, reports, reviews, and notifications. JWT-based authentication, role-aware administration, Socket.IO events, and carefully modelled swap states support a safer, more auditable marketplace experience.');
para(doc,'The implemented scope includes authentication, user profiles, listing management, discovery filters, wishlists, a structured swap lifecycle, delivery coordination, manual payment records, real-time chat foundations, trust and moderation tools, notifications, demo data, smoke checks, and deployment guidance. The report documents the problem, requirements, architecture, implementation, testing approach, deployment considerations, limitations, and future improvements.');
para(doc,'Keywords: clothing exchange, circular fashion, marketplace, React, Express, PostgreSQL, Prisma, JWT, Socket.IO.',True,10,TEAL,after=10)
heading(doc,'Table of Contents',1)
for item in ['1. Introduction','2. Problem Statement and Objectives','3. Requirement Analysis','4. Technology Stack','5. System Architecture and Design','6. Database Design','7. Module Implementation','8. User Workflow','9. Testing and Quality Assurance','10. Security and Deployment','11. Limitations and Future Scope','12. Conclusion','References']:
    para(doc,item,False,10,'202B36',after=3)
pagebreak(doc)

heading(doc,'1. Introduction',1); para(doc,'Fast fashion and short product lifecycles create an opportunity for people to exchange quality clothing that they no longer use. Existing informal exchanges often lack searchable listings, clear ownership, messaging, transaction states, and a way to handle delivery or disputes. SwapWear brings these steps into one responsive marketplace.');
heading(doc,'1.1 Project Overview',2); para(doc,'SwapWear allows registered users to publish clothing listings, explore items from other users, save favourites, propose item-for-item swaps, discuss a proposed exchange, select a delivery method, and confirm the exchange. Administrators can view moderation queues, reports, and marketplace statistics.');
heading(doc,'1.2 Project Scope',2)
for x in ['Public browsing of marketplace listings and item details.','Secure account registration, login, logout, and password reset flow.','Listing creation, editing, media support, search, filters, and wishlist management.','Swap initiation, acceptance, locking, cancellation, completion, disputes, and event history.','Delivery records, addresses, tracking/proof fields, manual payments, chat, calls/signalling, notifications, reviews, reports, and administration.'] : bullet(doc,x)
heading(doc,'2. Problem Statement and Objectives',1); para(doc,'The central problem is that clothing exchange is usually fragmented across informal chats or generic classified platforms. Users need a focused system that represents both items in a proposed exchange, prevents conflicting commitments, and makes the status of each transaction understandable.');
table(doc,['Objective','How SwapWear addresses it'],[['Enable sustainable reuse','Makes pre-owned clothing discoverable and reusable through a dedicated swap marketplace.'],['Reduce friction','Provides listings, search, wishlists, requests, chat, and delivery coordination in one flow.'],['Protect transaction integrity','Uses explicit swap statuses, listing reservation, confirmations, events, reports, and moderation.'],['Support administration','Provides admin-only routes and moderation/trust services for reports, payments, users, and listings.']], [2.1,4.4])
pagebreak(doc)

heading(doc,'3. Requirement Analysis',1); heading(doc,'3.1 Functional Requirements',2)
table(doc,['ID','Requirement'],[['FR-01','Users shall be able to register, log in, log out, and reset a password.'],['FR-02','Users shall be able to create, view, update, and manage clothing listings.'],['FR-03','Users shall be able to search, filter, sort, and save listings to a wishlist.'],['FR-04','Users shall be able to create and manage swap requests involving two listings.'],['FR-05','The system shall reserve items after an accepted swap and prevent conflicting acceptance.'],['FR-06','Users shall be able to choose local or courier delivery and record delivery progress.'],['FR-07','The platform shall support chat, notifications, reviews, reports, and role-based administration.']], [.65,5.85])
heading(doc,'3.2 Non-functional Requirements',2)
for x in ['Responsive and accessible interface for mobile, tablet, and desktop layouts.','Role-based access control for protected user and administrator pages.','Consistent API error handling, offline feedback, loading states, and empty states.','Relational integrity and lifecycle safety for swaps, delivery records, reports, and payments.','Environment-based configuration for local development and production deployment.'] : bullet(doc,x)
heading(doc,'3.3 Actors',2); table(doc,['Actor','Primary responsibilities'],[['Visitor','Browse public pages, explore listings, and access authentication.'],['Registered user','Manage profile, listings, wishlists, requests, chats, deliveries, reviews, and settings.'],['Administrator','Monitor users/listings, resolve reports, review payment queue, and manage marketplace trust.']], [1.65,4.85])
pagebreak(doc)

heading(doc,'4. Technology Stack',1); table(doc,['Layer','Technology','Purpose'],[['Client','React 19, Vite 8, React Router','Component-based responsive single-page application and route management.'],['Styling','Tailwind CSS, custom design system, Framer Motion','Consistent layout, theme, component styling, and animation.'],['Server','Node.js, Express 5','REST APIs, middleware, request validation, and business logic.'],['Database','PostgreSQL, Prisma 6','Relational persistence, schema modelling, constraints, and queries.'],['Authentication','JWT, bcryptjs, cookies','Credential verification, tokens, refresh cookies, and protected routes.'],['Realtime','Socket.IO','Conversation rooms, typing events, notifications, and call signalling.'],['Media','Multer / Cloudinary integration','Profile and listing media upload support.'],['Deployment','Vercel frontend configuration','Production hosting configuration for the client application.']], [1.05,1.8,3.65])
heading(doc,'5. System Architecture and Design',1); para(doc,'SwapWear follows a client-server architecture. The React client renders pages and calls backend APIs through service modules. Express routes delegate to controller and service layers, which use Prisma to access PostgreSQL. Socket.IO augments the API with event-based communication for chats, notifications, and WebRTC signalling.');
table(doc,['Layer','Main elements','Responsibility'],[['Presentation','Pages, components, hooks, contexts','Render responsive UI, capture user interaction, maintain client state.'],['Application','Services, routes, controllers','Apply validation, authorization, lifecycle rules, and API responses.'],['Persistence','Prisma schema and PostgreSQL','Store related marketplace entities with indexes and constraints.'],['Realtime','Socket.IO rooms and events','Deliver conversation, typing, notification, and call-signalling events.']], [1.2,2.6,2.7])
heading(doc,'5.1 High-level Data Flow',2); para(doc,'User action → React page/component → frontend service → Express route/controller → domain service → Prisma → PostgreSQL → JSON response → interface update. For real-time interactions: Socket.IO client ↔ conversation/user room ↔ Socket.IO server.');
pagebreak(doc)

heading(doc,'6. Database Design',1); para(doc,'PostgreSQL is used because the project requires strong relationships and transactional consistency among users, listings, swaps, delivery records, payments, moderation, and communication. Prisma expresses the schema and creates a type-aware interface for database operations.');
table(doc,['Entity','Key relationships / purpose'],[['User and Profile','A user has one optional profile and may own listings, swaps, messages, notifications, payments, and reviews.'],['Listing and WishlistItem','A listing belongs to a user; wishlist items form a unique user-listing relationship.'],['Swap and SwapEvent','A swap links requester and owner users plus one listing from each party; events record lifecycle actions.'],['DeliveryOrder and Address','Delivery orders belong to swaps and can reference pickup/drop addresses.'],['ChatConversation and ChatMessage','Conversation may be linked to a swap and contains message/call history.'],['Report, Review, Payment','Trust, moderation, review, and payment entities retain auditable transaction information.']], [1.85,4.65])
heading(doc,'6.1 Core Swap States',2); table(doc,['Status','Meaning'],[['PENDING','A proposal awaits owner action.'],['ACCEPTED','Both listings are reserved for the swap.'],['REJECTED / CANCELLED / EXPIRED','The exchange did not proceed; eligible listings can return to availability.'],['SHIPPED / DELIVERED','Delivery has progressed.'],['COMPLETED','The exchange has been successfully completed and may be archived.'],['DISPUTED / FAILED','A conflict or exception requires resolution.']], [1.9,4.6])
pagebreak(doc)

heading(doc,'7. Module Implementation',1); table(doc,['Module','Implemented capabilities'],[['Authentication & profiles','Email/password sign-up and login, JWT flow, refresh cookies, password reset, avatar/profile management, role-aware access.'],['Marketplace','Home, explore, item detail, categories, product filters, media galleries, related items, dashboard, and responsive UI.'],['Listings & wishlist','Create/edit listings, images, statuses, public visibility, search/filter/sort, and unique wishlist items.'],['Swap management','Request, accept, reject, cancel, complete, archive, confirmation, dispute, item reservation, and lifecycle events.'],['Delivery & payments','Address book, local/courier methods, two delivery legs, tracking/proof fields, manual payment records, status updates.'],['Chat & notifications','Conversations, messages, uploads, typing events, conversation rooms, notification list/read/delete actions, call session signalling.'],['Trust & admin','Marketplace reports, review records, report resolution/blocking, user/listing moderation, payment and dashboard queues.']], [1.65,4.85])
heading(doc,'7.1 Frontend Design',2); para(doc,'The frontend uses route-level lazy loading and Suspense to reduce initial load work. Public routes include home, explore, item details, and community. Protected routes cover dashboard, profiles, wishlist, settings, listings, swaps, chat, and administration. Reusable UI, layout, product, swap, chat, auth, and admin components keep concerns separated.');
heading(doc,'7.2 Backend Design',2); para(doc,'The Express server configures Helmet, CORS, rate limiting, JSON parsing, cookies, upload access, health checks, domain routes, and centralized error handling. It also creates a Socket.IO server for user and conversation rooms. The API is organized around auth, users, listings, wishlist, swaps, chat, notifications, delivery, payments, and trust.');
pagebreak(doc)

heading(doc,'8. User Workflow',1); table(doc,['Step','User journey'],[['1. Discover','A visitor browses the home page or explores public listings.'],['2. Authenticate','The user creates an account or logs in to access protected actions.'],['3. List or select','The user creates a listing or selects an existing item to exchange.'],['4. Propose swap','The requester chooses an offered item and submits a swap request with a message.'],['5. Decide and lock','The owner accepts or rejects. On acceptance, the system reserves both listings and records the event.'],['6. Coordinate','Participants use the deal room/chat, choose delivery, and update tracking or handover information.'],['7. Confirm','Participants confirm exchange/delivery; the swap progresses to completion or a dispute is opened.'],['8. Review and moderate','Users can leave reviews or reports; administrators can resolve moderation issues.']], [.65,5.85])
heading(doc,'8.1 Important Business Rule',2); para(doc,'When a swap is accepted, the application reserves both involved listings and expires competing pending requests. If a swap is rejected, cancelled, or fails, the design supports relisting and reviving eligible requests. This is the key rule that prevents a single item from being committed to multiple accepted swaps.');
heading(doc,'9. Testing and Quality Assurance',1); para(doc,'The repository includes runnable checks and a repeatable demo workflow. The demo seed provides users, profiles, addresses, listings, a pending swap, a conversation, notifications, and an admin report so functionality can be tested without starting with an empty database.');
table(doc,['Check','Purpose'],[['npm run build','Build the frontend for production and identify compilation issues.'],['npm run lint','Run ESLint checks over the frontend codebase.'],['npx prisma validate','Validate the Prisma database schema.'],['npm run seed:demo','Populate idempotent demo users and marketplace data.'],['npm run smoke','Confirm environment loading, database reachability, and basic core-table reads.'],['swap lifecycle check','Exercise lifecycle-focused validation for swap rules.']], [2.0,4.5])
pagebreak(doc)

heading(doc,'10. Security and Deployment',1); heading(doc,'10.1 Security Measures',2)
for x in ['Passwords are hashed with bcryptjs; authenticated API access uses JWT tokens and refresh cookies.','Protected and administrator-only frontend routes prevent unauthorized page access; backend middleware enforces API-side authorization.','Helmet applies security-oriented HTTP headers, CORS restricts configured client origins, and rate limiting reduces abusive request volume.','The Prisma schema uses relations, unique constraints, indexes, and enumerated states to support data consistency.','Production configuration validates required environment values, while secrets remain in environment files rather than source control.'] : bullet(doc,x)
heading(doc,'10.2 Deployment',2); para(doc,'The supplied frontend deployment URL is https://swapwear-iota.vercel.app/. The repository includes Vercel configuration for the frontend. A complete production deployment also requires a hosted PostgreSQL database, backend hosting, configured CORS origins, secure JWT secrets, persistent object storage for uploads, and optional payment/courier provider credentials.');
heading(doc,'10.3 Operational Checklist',2); table(doc,['Area','Production action'],[['Database','Provision PostgreSQL, apply schema deployment, validate health endpoint.'],['Secrets','Set database URL, JWT secrets, client URL, OAuth callback base, and provider credentials.'],['Media','Use persistent cloud storage instead of local upload storage.'],['Monitoring','Add logs, error tracking, uptime checks, and backup strategy.'],['Manual QA','Verify protected flows, media, swaps, deliveries, moderation, and mobile layout.']], [1.45,5.05])
heading(doc,'11. Limitations and Future Scope',1); para(doc,'The project includes production-ready foundations but some integrations are intentionally represented by manual or adapter-ready workflows. The following work would strengthen a public launch.');
for x in ['Integrate verified payment gateways such as Razorpay or Stripe and process webhooks.','Connect courier providers for labels, live tracking, and automated shipping status.','Complete WebRTC peer-media UI and configure TURN servers for reliable audio/video calls.','Move all uploads to durable object storage with CDN delivery and content scanning.','Add automated unit, integration, end-to-end, accessibility, and load tests in CI/CD.','Introduce recommendation ranking, saved searches, richer analytics, and image-quality moderation.'] : bullet(doc,x)
pagebreak(doc)

heading(doc,'12. Conclusion',1); para(doc,'SwapWear demonstrates a complete approach to a modern circular-fashion marketplace. The project moves beyond a simple item catalogue by modelling the operational details that make an exchange credible: identity, listings, mutually linked items, reservations, state transitions, delivery, communication, reviews, reports, and administration.');
para(doc,'Its implementation uses a contemporary, maintainable stack: React and Vite on the client; Express and Socket.IO on the server; PostgreSQL and Prisma for reliable relational data; and JWT-based access control. The documented QA workflow, demo data, production configuration, and future-integration plan give the project a practical path from academic demonstration to a deployable marketplace.');
heading(doc,'References',1)
for x in ['SwapWear source repository. https://github.com/SatyamKushwaha3232/Swapwear','SwapWear live deployment. https://swapwear-iota.vercel.app/','React Documentation. https://react.dev/','Vite Documentation. https://vite.dev/','Express Documentation. https://expressjs.com/','Prisma Documentation. https://www.prisma.io/docs/','PostgreSQL Documentation. https://www.postgresql.org/docs/','Socket.IO Documentation. https://socket.io/docs/'] : para(doc,x,False,9.5,TEAL,after=4)
para(doc,'Appendix A - Repository structure: frontend/ contains pages, components, hooks, services, contexts, assets, and design-system files. backend/ contains configuration, routes, controllers, domain modules, services, middleware, Prisma schema, scripts, and environment templates.',False,9.5,GRAY,after=8,before=16)

doc.core_properties.title='SwapWear Detailed Project Report'; doc.core_properties.author='Satyam Kushwaha'; doc.core_properties.subject='Full-stack clothing swap marketplace project report'
doc.save(FILE)
print(FILE.resolve())
